#!/usr/bin/python
# encoding=utf8
from docx import *
import sys
import os
import re
os.environ.setdefault('DJANGO_SETTINGS_MODULE', 'ants.settings')
from populate import populate, add_ant, add_species

document = Document('antdata.docx')

def parse_cell(text):

	main_dict = {"country":"",
	             "state":"",
	             "landmark":"",
	             "collection":"",
	             "date":"",
	             "owner":"",
	             "description":"",
	             "GPS_lat":"",
	             "GPS_long":"",
	             "ID_date":"",
	             "discovery":"",
	             "aID":"",
	             "specimen":"",
	             "family":"",
	             "genus":"",
	             "species":"",
	             "specimen_ID":"",}
	

	#get location and everything in Cell #1	
	location = re.search(r'.*?United-States|.*?Canada', text)
	if location:
		adress_data = location.group().split(',')
		adress_data.reverse()	
		if len(adress_data) == 1:
			print (u"Country: {0}".format(adress_data[0]).encode('UTF-8'))
			main_dict["country"]=adress_data[0].encode('UTF-8')
		elif len(adress_data) == 2:
			print (u"Country: {0}".format(adress_data[0]).encode('UTF-8'))
			main_dict["country"]=adress_data[0].encode('UTF-8')
			print (u"State: {0}".format(adress_data[1]).encode('UTF-8'))
			main_dict["state"]=adress_data[1].encode('UTF-8')
		elif len(adress_data) > 2:
			print (u"Country: {0}".format(adress_data[0]).encode('UTF-8'))
			main_dict["country"]=adress_data[0].encode('UTF-8')
                        print (u"State: {0}".format(adress_data[1]).encode('UTF-8'))
			main_dict["state"]=adress_data[1].encode('UTF-8')
			i=2
			rest_of_adress=""
			while i < len(adress_data):
				rest_of_adress = rest_of_adress + " " + adress_data[i]
				i=i+1
                        print (u"Also: {0}".format(rest_of_adress.strip()).encode('UTF-8'))
			main_dict["landmark"] = rest_of_adress.strip().encode('UTF-8')
		lab = re.search(r'\(Lab.*?\)', text)
		if lab:
			lab=lab.group(0)
			print (u"Specimen from: {0}".format(lab).encode('UTF-8'))
			main_dict["collection"]=lab.encode('UTF-8')
		date = re.search(r'.*?-20..|.*?-19..', text)
                if date:
                        date = date.group(0)
			print (u"Date: {0}".format(date).encode('UTf-8'))
			main_dict["date"]=date.encode('UTF-8')

		
		text_split_list = text.split('\n') 

		owner =  text_split_list[len(text_split_list)-1]
		
		print (u"Owner: {0}".format(owner).encode('UTF-8'))
		main_dict["owner"]=owner.encode('UTF-8')



	#get ID and everything in Cell #2
	ID = re.search(r'ID\: .*', text)
	if ID:
		ID = ID.group().strip("ID:")
		
		text2_split_list = text.split('\n')
		description = text2_split_list[0]
		description_date = text2_split_list[len(text2_split_list)-1]
		print "HI I EXIST".encode('UTF-8')
		print (u"Description: {0}".format(description).encode('UTF-8'))
		print (u"ID: {0}".format(ID).encode('UTF-8'))
		print (u"Edited on: {0}".format(description_date).encode('UTF-8'))
		main_dict["description"]=description.encode('UTF-8')	
		main_dict["aID"]=ID.encode('UTF-8')
		main_dict["ID_date"]=description_date.encode('UTF-8')


		if len(text2_split_list) <= 3:

			print (u"MIDDLE: {0}".format(text2_split_list[1]).encode('UTF-8'))			
		else:
			regex_lat=re.compile("±\s*[NS] .*[’']".decode('UTF-8'))
			GPS_lat = re.search(regex_lat, text)
			regex_long = re.compile("WO .*[’']".decode('UTF-8'))
			GPS_long = re.search(regex_long, text)
			if GPS_lat and GPS_long:
				GPS_lat = GPS_lat.group()
				GPS_long = GPS_long.group()
				#Format function acts weird here, so I concatenate the old fashion way
				print ("GPS: " + (GPS_lat).encode('UTF-8') + "   " + (GPS_long).encode('UTF-8') )
				main_dict["GPS_lat"]=GPS_lat.encode('UTF-8')
				main_dict["GPS_long"]=GPS_long.encode('UTF-8')
	#get discovery date and everything in cell #3
	
	discovery = re.search(r'\(.* [1-2][0-9][0-9][0-9]\)', text)
	if discovery:
		discovery = discovery.group().strip('()')
		text3_split_list = text.split('\n')
		genus = text3_split_list[1].split(' ')[0]
		species = text3_split_list[1].split(' ')[1]	


		print (u"Discovery: {0}".format(discovery).encode('UTF-8'))
		print (u"FFamily: {0}".format(text3_split_list[0]).encode('UTF-8'))
		print (u"Genus: {0}".format(genus).encode('UTF-8'))
		print (u"Species: {0}".format(species).encode('UTF-8'))
		main_dict["discovery"]=discovery.encode('UTF-8')		
		main_dict["family"]=text3_split_list[0].encode('UTF-8')
		main_dict["genus"]=genus.encode('UTF-8')
		main_dict["species"]=species.encode('UTF-8')


		if len(text3_split_list) == 5:
			print (u"Code: {0}".format(text3_split_list[4]).encode('UTF-8'))
			main_dict["specimen_ID"]=text3_split_list[4].encode('UTF-8')
		else:
			print (u"Code: {0}".format(text3_split_list[3]).encode('UTF-8'))
			main_dict["specimen_ID"]=text3_split_list[3].encode('UTF-8')

	
	return main_dict



tables = document.tables
i = 0
for table in tables:

	x = 0
	y = 0
	p = 0
	cell_counter = 0
	main_dict={}
#document.tables[i].cell(x,y).paragraphs[p].text != None



	while True:
		try:
			test_end = document.tables[i].cell(x,y).paragraphs[p].text
			p=0	
			while True:

				try:
					paragraphs = document.tables[i].cell(x,y).paragraphs
					
					text=""
					for paragraph in paragraphs:
						text = text +'\n'+ paragraph.text
						text = text.strip('\n')
						#print text.encode('UTF-8')
					#	print (u"paragraph {0}:   {1}".format(p, document.tables[i].cell(x,y).paragraphs[p].text).encode('UTF-8'))
						p = p+1
					print text.encode('UTF-8')
					main_dict = dict(main_dict.items() + parse_cell(text).items())
					cell_counter = cell_counter + 1
					if cell_counter == 3:
						populate(main_dict)
						cell_counter = 0
						main_dict={}
					x=x+1
					print ("----------------y+1------------------")
					p=0
				except IndexError:
					x=0
					break
			

		except IndexError:
			break
		p=0
		y = y + 1
		print ("-----------------------x+1----------------------------")
	i = i + 1

print ("REACHED").encode('UTF-8')
